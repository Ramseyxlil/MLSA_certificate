import os
import pandas as pd
import convertapi
from flask import Flask, request, send_file, render_template, redirect, url_for
from docx import Document
from docx.shared import Pt, RGBColor
from zipfile import ZipFile

app = Flask(__name__)

# Define the path for uploads and generated files
UPLOAD_FOLDER = 'uploads'
CERTIFICATE_FOLDER = 'certificates'
ZIP_FOLDER = 'zips'
TEMPLATE_FILE = 'certificate_template.docx'

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

if not os.path.exists(CERTIFICATE_FOLDER):
    os.makedirs(CERTIFICATE_FOLDER)

if not os.path.exists(ZIP_FOLDER):
    os.makedirs(ZIP_FOLDER)

# ConvertAPI secret key
convertapi.api_secret = 'yQBA8IqyogHySlgp'


@app.route('/')
def upload_form():
    return render_template('upload.html')


@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return 'No file part'
    file = request.files['file']
    if file.filename == '':
        return 'No selected file'
    if file:
        event_name = request.form['event_name']
        ambassador_name = request.form['ambassador_name']
        file_path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(file_path)

        # Process the file and generate certificates
        generate_certificates(file_path, event_name, ambassador_name)

        zip_file_path = create_zip(event_name)
        download_url = url_for('download_file', path=zip_file_path)
        return redirect(f'/?download_url={download_url}')


@app.route('/download/<path:path>')
def download_file(path):
    return send_file(path, as_attachment=True)


def apply_font_style(run, font_size, color, bold):
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)


def generate_certificate(participant_name, event_name, ambassador_name):
    doc = Document(TEMPLATE_FILE)
    for p in doc.paragraphs:
        if '{PARTICIPANT_NAME}' in p.text:
            # Replace and apply font style for participant name
            p.text = p.text.replace('{PARTICIPANT_NAME}', participant_name)
            for run in p.runs:
                apply_font_style(run, font_size=24, color=(0, 0, 255), bold=True)
        if '{EVENT_NAME}' in p.text:
            # Replace and apply font style for event name
            p.text = p.text.replace('{EVENT_NAME}', event_name)
            for run in p.runs:
                apply_font_style(run, font_size=14, color=(0, 0, 0), bold=True)
        if '{AMBASSADOR_NAME}' in p.text:
            # Replace and apply font style for ambassador name
            p.text = p.text.replace('{AMBASSADOR_NAME}', ambassador_name)
            for run in p.runs:
                apply_font_style(run, font_size=11, color=(0, 0, 0), bold=True)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if '{PARTICIPANT_NAME}' in p.text:
                        # Replace and apply font style for participant name
                        p.text = p.text.replace('{PARTICIPANT_NAME}', participant_name)
                        for run in p.runs:
                            apply_font_style(run, font_size=24, color=(0, 0, 255), bold=True)
                    if '{EVENT_NAME}' in p.text:
                        # Replace and apply font style for event name
                        p.text = p.text.replace('{EVENT_NAME}', event_name)
                        for run in p.runs:
                            apply_font_style(run, font_size=14, color=(0, 0, 0), bold=True)
                    if '{AMBASSADOR_NAME}' in p.text:
                        # Replace and apply font style for ambassador name
                        p.text = p.text.replace('{AMBASSADOR_NAME}', ambassador_name)
                        for run in p.runs:
                            apply_font_style(run, font_size=11, color=(0, 0, 0), bold=True)

    docx_path = os.path.join(CERTIFICATE_FOLDER, f'{participant_name}.docx')
    doc.save(docx_path)
    return docx_path


def generate_certificates(file_path, event_name, ambassador_name):
    df = pd.read_csv(file_path)
    for index, row in df.iterrows():
        participant_name = row['Name']
        docx_path = generate_certificate(participant_name, event_name, ambassador_name)
        convert_to_pdf(docx_path)


def convert_to_pdf(docx_path):
    result = convertapi.convert('pdf', {
        'File': docx_path
    }, from_format='docx')

    pdf_path = docx_path.replace('.docx', '.pdf')
    result.file.save(pdf_path)
    os.remove(docx_path)




def create_zip(event_name):
    zip_file_path = os.path.join(ZIP_FOLDER, f'{event_name}_certificates.zip')
    with ZipFile(zip_file_path, 'w') as zipf:
        for folder_name, subfolders, filenames in os.walk(CERTIFICATE_FOLDER):
            for filename in filenames:
                if filename.endswith('.pdf'):
                    file_path = os.path.join(folder_name, filename)
                    zipf.write(file_path, os.path.basename(file_path))
    return zip_file_path


if __name__ == '__main__':
    app.run(debug=True)
